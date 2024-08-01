import openai
import tiktoken
import streamlit as st
from docx import Document
import os
import shutil
from io import BytesIO

# Ensure necessary directories exist
os.makedirs('uploads', exist_ok=True)
os.makedirs('output', exist_ok=True)

# Extract OpenAI API key from environment variable
openai_api_key = st.secrets["OPENAI_API_KEY"]
openai.api_key = openai_api_key

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def clear_directory(dir_path):
    if os.path.exists(dir_path):
        shutil.rmtree(dir_path)
    os.makedirs(dir_path, exist_ok=True)

def count_tokens(text, model="gpt-3.5-turbo"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

def latex_to_readable(latex_code):
    combined_prompt = f"""
    You are an intelligent assistant. Your task is to convert LaTeX code and programming code in the given text to plain English and provide summaries for tables and images. Follow these specific instructions:

    1. Keep the plain text exactly as it is. Do not provide any overview or summary of the entire text.
    2. Convert any LaTeX code into a human-readable format. For example, $\\alpha$ should be read as "alpha", and $ax^2+bx+c=0$ should be read as "a x squared plus b x plus c equals zero".
    3. For programming code, read the code line by line and provide a high-level description of what the code is doing. For example, if the code is "for i in range(10): print(i)", you can say "A loop that prints numbers from 0 to 9".
    4. For tables, provide a summarized explanation of the concept covered in the table without mentioning the formatting. For example, if a table shows a logical representation of inputs and outputs, explain the logical relationship in plain English.
    5. For images, mention that there is an image at that position, but do not narrate the source or provide additional context about the image.

    Text to convert:
    {latex_code}
    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an intelligent assistant."},
                {"role": "user", "content": combined_prompt}
            ]
        )
        human_readable_text = response.choices[0].message['content'].strip()

        # Calculate tokens
        prompt_tokens = count_tokens(combined_prompt)
        response_tokens = count_tokens(human_readable_text)
        total_tokens = prompt_tokens + response_tokens

        return human_readable_text, total_tokens

    except Exception as e:
        return f"Error: {e}", 0

def save_text_to_docx(text, output_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

# Streamlit app
st.title("LaTeX and Code to Human-Readable Text Converter")

st.write("""Upload a DOCX file with LaTeX content and/or programming code, and get a human-readable text version.""")

uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")

if uploaded_file is not None:
    clear_directory('uploads')
    clear_directory('output')
    
    # Save uploaded file
    file_path = os.path.join('uploads', uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    text = extract_text_from_docx(file_path)

    # Use ChatGPT Turbo to convert text
    try:
        with st.spinner("Converting LaTeX and code to human-readable text..."):
            converted_text, total_tokens = latex_to_readable(text)
        st.success("Conversion successful!")
    except Exception as e:
        st.error(f"Error communicating with OpenAI: {e}")

    st.write("### Converted Text")
    st.write(converted_text)

    st.write("### Total Tokens Used")
    st.write(total_tokens)

    # Save converted text to a DOCX file
    docx_output_path = os.path.join('output', 'converted_text.docx')
    save_text_to_docx(converted_text, docx_output_path)

    # Provide a download link for the DOCX file
    with open(docx_output_path, "rb") as f:
        st.download_button(
            label="Download DOCX",
            data=f,
            file_name="converted_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
